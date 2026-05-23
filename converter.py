import io
import json
from pathlib import Path
from copy import deepcopy

DEFAULT_TEMPLATE = {
    "document": {
        "default_font": "Times New Roman",
        "default_size": 12,
        "page_margins": {"top": 2.54, "bottom": 2.54, "left": 3.18, "right": 3.18},
    },
    "styles": {
        "heading_1": {"font_name": "Arial", "font_size": 24, "bold": True, "alignment": "center", "color": "#1a1a1a"},
        "heading_2": {"font_name": "Arial", "font_size": 18, "bold": True, "color": "#2d2d2d"},
        "heading_3": {"font_name": "Arial", "font_size": 14, "bold": True, "color": "#3d3d3d"},
        "heading_4": {"font_name": "Arial", "font_size": 12, "bold": True},
        "heading_5": {"font_name": "Arial", "font_size": 11, "bold": True},
        "heading_6": {"font_name": "Arial", "font_size": 10, "bold": True},
        "paragraph": {"font_name": "Times New Roman", "font_size": 12},
        "bold": {"bold": True},
        "italic": {"italic": True},
        "inline_code": {"font_name": "Courier New", "font_size": 10.5, "bg_color": "#F0F0F0"},
        "code_block": {"font_name": "Courier New", "font_size": 10, "bg_color": "#F0F0F0", "border_color": "#CCCCCC", "border_width_pt": 1},
        "link": {"color": "#0000EE", "underline": True},
        "table_header": {"font_name": "Arial", "font_size": 10, "bold": True},
        "table_cell": {"font_name": "Arial", "font_size": 10},
        "list_bullet": {"indent_cm": 1.27},
        "list_number": {"indent_cm": 1.27},
        "image": {"alignment": "center"},
    },
}


class ConversionError(Exception):
    pass


def _hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def _apply_run_style(run, style):
    if style.get("bold"):
        run.bold = True
    if style.get("italic"):
        run.italic = True
    if "font_name" in style:
        run.font.name = style["font_name"]
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), style["font_name"])
    if "font_size" in style:
        run.font.size = docx_shared.Pt(style["font_size"])
    if "color" in style:
        r, g, b = _hex_to_rgb(style["color"])
        run.font.color.rgb = docx_shared.RGBColor(r, g, b)
    if style.get("underline"):
        run.underline = True


def _apply_paragraph_shading(paragraph, bg_color):
    from lxml import etree

    r, g, b = _hex_to_rgb(bg_color)
    shading_elm = etree.SubElement(paragraph._p.get_or_add_pPr(), docx_oxml.ns.qn("w:shd"))
    shading_elm.set(docx_oxml.ns.qn("w:val"), "clear")
    shading_elm.set(docx_oxml.ns.qn("w:color"), "auto")
    shading_elm.set(docx_oxml.ns.qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")


def _apply_paragraph_border(paragraph, border_color, border_width_pt):
    from lxml import etree

    r, g, b = _hex_to_rgb(border_color)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, docx_oxml.ns.qn("w:pBdr"))
    for side in ("top", "left", "bottom", "right"):
        border = etree.SubElement(pBdr, docx_oxml.ns.qn(f"w:{side}"))
        border.set(docx_oxml.ns.qn("w:val"), "single")
        border.set(docx_oxml.ns.qn("w:sz"), str(int(border_width_pt * 8)))
        border.set(docx_oxml.ns.qn("w:space"), "4")
        border.set(docx_oxml.ns.qn("w:color"), f"{r:02X}{g:02X}{b:02X}")


import markdown_it
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import oxml as docx_oxml
from docx import shared as docx_shared


md_parser = markdown_it.MarkdownIt("commonmark", {"typographer": True}).enable(["table", "strikethrough"])


def parse_markdown(md_text):
    if not isinstance(md_text, str):
        raise ConversionError(f"输入必须是字符串，收到 {type(md_text).__name__}")
    return md_parser.parse(md_text)


def build_ir(tokens, depth=0):
    if depth > 10:
        raise ConversionError("Markdown 嵌套层级过深")
    ir = []
    i = 0
    while i < len(tokens):
        token = tokens[i]
        t_type = token.type

        if t_type == "heading_open":
            level = int(token.tag[1])
            inline_tokens = tokens[i + 1]
            text = _extract_text(inline_tokens) if inline_tokens.children else ""
            plain = _extract_plain_text(inline_tokens) if inline_tokens.children else ""
            ir.append(
                {
                    "type": "heading",
                    "level": level,
                    "text": text,
                    "plain": plain,
                    "inline": _build_inline_ir(inline_tokens.children) if inline_tokens.children else [],
                }
            )
            i += 3  # skip heading_open, inline, heading_close

        elif t_type == "paragraph_open":
            inline = tokens[i + 1]
            text = _extract_text(inline) if inline.children else ""
            ir.append(
                {
                    "type": "paragraph",
                    "text": text,
                    "inline": _build_inline_ir(inline.children) if inline.children else [],
                }
            )
            i += 3  # skip paragraph_open, inline, paragraph_close

        elif t_type == "bullet_list_open":
            items = _build_list_ir(tokens, i, "bullet")
            ir.append({"type": "list", "list_type": "bullet", "items": items})
            i = _skip_list(tokens, i)

        elif t_type == "ordered_list_open":
            items = _build_list_ir(tokens, i, "ordered")
            ir.append({"type": "list", "list_type": "ordered", "items": items})
            i = _skip_list(tokens, i)

        elif t_type == "fence":
            ir.append(
                {
                    "type": "code_block",
                    "language": token.info.strip() if token.info else "",
                    "content": token.content,
                }
            )
            i += 1

        elif t_type == "table_open":
            table_data = _build_table_ir(tokens, i)
            ir.append({"type": "table", "headers": table_data["headers"], "rows": table_data["rows"]})
            i = table_data["end"]

        elif t_type == "blockquote_open":
            inner = _build_blockquote_ir(tokens, i)
            ir.append({"type": "blockquote", "content": inner})
            i = _skip_blockquote(tokens, i)

        elif t_type == "hr":
            ir.append({"type": "hr"})
            i += 1

        else:
            i += 1

    return ir


def _extract_text(inline_token):
    parts = []
    if hasattr(inline_token, "children") and inline_token.children:
        for child in inline_token.children:
            if child.type == "text":
                parts.append(child.content)
            elif child.type == "softbreak":
                parts.append(" ")
            elif child.type == "hardbreak":
                parts.append("\n")
            elif hasattr(child, "children") and child.children:
                parts.append(_extract_text(child))
    elif hasattr(inline_token, "content"):
        parts.append(inline_token.content)
    return "".join(parts)


def _extract_plain_text(inline_token):
    if hasattr(inline_token, "children") and inline_token.children:
        parts = []
        for child in inline_token.children:
            if child.type == "text":
                parts.append(child.content)
            elif hasattr(child, "children") and child.children:
                parts.append(_extract_plain_text(child))
        return "".join(parts)
    return ""


def _build_inline_ir(children):
    result = []
    for child in children:
        if child.type == "text":
            result.append({"type": "text", "content": child.content})
        elif child.type == "strong_open":
            inner = _build_inline_ir(child.children) if child.children else []
            result.append({"type": "bold", "content": inner})
        elif child.type == "em_open":
            inner = _build_inline_ir(child.children) if child.children else []
            result.append({"type": "italic", "content": inner})
        elif child.type == "code_inline":
            result.append({"type": "inline_code", "content": child.content})
        elif child.type == "link_open":
            inner = _build_inline_ir(child.children) if child.children else []
            href = child.attrs.get("href", "")
            result.append({"type": "link", "href": href, "content": inner})
        elif child.type == "image":
            src = child.attrs.get("src", "")
            alt = child.content if hasattr(child, "content") and child.content else ""
            result.append({"type": "image", "src": src, "alt": alt})
        elif child.type == "softbreak":
            result.append({"type": "text", "content": " "})
        elif child.type == "hardbreak":
            result.append({"type": "hardbreak"})
        elif child.type == "s_open":
            inner = _build_inline_ir(child.children) if child.children else []
            result.append({"type": "strikethrough", "content": inner})
    return result


def _build_list_ir(tokens, start_idx, list_type):
    items = []
    i = start_idx + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "list_item_open":
            i += 1
            item_parts = []
            while i < len(tokens) and tokens[i].type not in ("list_item_close", "bullet_list_close", "ordered_list_close"):
                if tokens[i].type == "paragraph_open":
                    inline = tokens[i + 1]
                    text = _extract_text(inline) if inline.children else ""
                    item_parts.append(
                        {"type": "paragraph", "text": text, "inline": _build_inline_ir(inline.children) if inline.children else []}
                    )
                    i += 2
                elif tokens[i].type in ("bullet_list_open", "ordered_list_open"):
                    sub_type = "bullet" if tokens[i].type == "bullet_list_open" else "ordered"
                    sub_items = _build_list_ir(tokens, i, sub_type)
                    item_parts.append({"type": "list", "list_type": sub_type, "items": sub_items})
                    i = _skip_list(tokens, i)
                    continue
                i += 1
            items.append({"content": item_parts})
        elif token.type in ("bullet_list_close", "ordered_list_close"):
            break
        i += 1
    return items


def _skip_list(tokens, start_idx):
    depth = 0
    i = start_idx
    open_type = tokens[start_idx].type
    close_type = open_type.replace("_open", "_close")
    while i < len(tokens):
        if tokens[i].type == open_type:
            depth += 1
        elif tokens[i].type == close_type:
            depth -= 1
            if depth == 0:
                return i + 1
        i += 1
    return len(tokens)


def _build_table_ir(tokens, start_idx):
    headers = []
    rows = []
    i = start_idx + 1
    in_thead = False
    in_tbody = False

    while i < len(tokens):
        token = tokens[i]
        if token.type == "thead_open":
            in_thead = True
            i += 1
            continue
        elif token.type == "thead_close":
            in_thead = False
            i += 1
            continue
        elif token.type == "tbody_open":
            in_tbody = True
            i += 1
            continue
        elif token.type == "tbody_close":
            in_tbody = False
            i += 1
            continue
        elif token.type == "tr_open":
            row, end_i = _build_table_row(tokens, i + 1)
            i = end_i
            if in_thead:
                headers = row
            else:
                rows.append(row)
            continue
        elif token.type == "table_close":
            return {"headers": headers, "rows": rows, "end": i + 1}
        i += 1

    return {"headers": headers, "rows": rows, "end": len(tokens)}


def _build_table_row(tokens, start_idx):
    cells = []
    i = start_idx
    while i < len(tokens):
        token = tokens[i]
        if token.type == "td_open" or token.type == "th_open":
            inline = tokens[i + 1]
            text = _extract_text(inline) if inline.children else ""
            cells.append(
                {
                    "text": text,
                    "inline": _build_inline_ir(inline.children) if inline.children else [],
                }
            )
            i += 2
        elif token.type == "tr_close":
            return cells, i
        i += 1
    return cells, len(tokens)


def _build_blockquote_ir(tokens, start_idx):
    inner = []
    i = start_idx + 1
    depth = 1
    while i < len(tokens) and depth > 0:
        token = tokens[i]
        if token.type == "blockquote_open":
            depth += 1
        elif token.type == "blockquote_close":
            depth -= 1
            if depth == 0:
                break
        elif token.type == "paragraph_open":
            inline = tokens[i + 1]
            text = _extract_text(inline) if inline.children else ""
            inner.append(
                {
                    "type": "paragraph",
                    "text": text,
                    "inline": _build_inline_ir(inline.children) if inline.children else [],
                }
            )
            i += 2
        i += 1
    return inner


def _skip_blockquote(tokens, start_idx):
    depth = 0
    i = start_idx
    while i < len(tokens):
        if tokens[i].type == "blockquote_open":
            depth += 1
        elif tokens[i].type == "blockquote_close":
            depth -= 1
            if depth == 0:
                return i + 1
        i += 1
    return len(tokens)


MAX_IR_NODES = 5000


def render_docx(ir, template):
    doc = Document()

    doc_margin = template.get("document", {}).get("page_margins", {})
    for section in doc.sections:
        section.top_margin = Cm(doc_margin.get("top", 2.54))
        section.bottom_margin = Cm(doc_margin.get("bottom", 2.54))
        section.left_margin = Cm(doc_margin.get("left", 3.18))
        section.right_margin = Cm(doc_margin.get("right", 3.18))

    default_font = template.get("document", {}).get("default_font", "Times New Roman")
    default_size = template.get("document", {}).get("default_size", 12)
    style = doc.styles["Normal"]
    style.font.name = default_font
    style.font.size = Pt(default_size)

    _render_ir_nodes(doc, ir, template)
    return doc


def _render_ir_nodes(doc, ir, template, parent=None):
    node_count = 0
    for node in ir:
        node_count += 1
        if node_count > MAX_IR_NODES:
            break

        if node["type"] == "heading":
            _render_heading(doc, node, template)
        elif node["type"] == "paragraph":
            _render_paragraph(doc, node, template)
        elif node["type"] == "code_block":
            _render_code_block(doc, node, template)
        elif node["type"] == "table":
            _render_table(doc, node, template)
        elif node["type"] == "list":
            _render_list(doc, node, template)
        elif node["type"] == "blockquote":
            _render_blockquote(doc, node, template)
        elif node["type"] == "hr":
            doc.add_paragraph().add_run().add_break()
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            pBdr.append(bottom)
            pPr.append(pBdr)

    return node_count


def _render_heading(doc, node, template):
    level = min(node["level"], 6)
    style_key = f"heading_{level}"
    style_dict = template.get("styles", {}).get(style_key, {})

    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]

    if style_dict.get("alignment") == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style_dict.get("alignment") == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if node.get("inline"):
        _render_inline_runs(p, node["inline"], style_dict, template)
    else:
        run = p.add_run(node.get("plain", node.get("text", "")))
        _apply_run_style(run, style_dict)


def _render_paragraph(doc, node, template):
    style_dict = template.get("styles", {}).get("paragraph", {})
    p = doc.add_paragraph()
    if node.get("inline"):
        _render_inline_runs(p, node["inline"], style_dict, template)
    else:
        run = p.add_run(node.get("text", ""))
        _apply_run_style(run, style_dict)


def _render_inline_runs(paragraph, inline_items, base_style, template, link_style=None):
    for item in inline_items:
        if item["type"] == "text":
            run = paragraph.add_run(item["content"])
            _apply_run_style(run, base_style)
            if link_style:
                _apply_run_style(run, link_style)
        elif item["type"] == "bold":
            bold_style = {**base_style, **template.get("styles", {}).get("bold", {})}
            _render_inline_runs(paragraph, item.get("content", []), bold_style, template, link_style)
        elif item["type"] == "italic":
            italic_style = {**base_style, **template.get("styles", {}).get("italic", {})}
            _render_inline_runs(paragraph, item.get("content", []), italic_style, template, link_style)
        elif item["type"] == "strikethrough":
            s_style = {**base_style}
            inner_items = item.get("content", [])
            for inner in inner_items:
                if inner["type"] == "text":
                    run = paragraph.add_run(inner["content"])
                    _apply_run_style(run, s_style)
                    run.font.strike = True
                    if link_style:
                        _apply_run_style(run, link_style)
        elif item["type"] == "inline_code":
            code_style = {**base_style, **template.get("styles", {}).get("inline_code", {})}
            run = paragraph.add_run(item["content"])
            _apply_run_style(run, code_style)
            if link_style:
                _apply_run_style(run, link_style)
        elif item["type"] == "link":
            link_style_dict = {**base_style, **template.get("styles", {}).get("link", {})}
            href = item.get("href", "")
            _render_link(paragraph, item.get("content", []), base_style, link_style_dict, template, href)
        elif item["type"] == "image":
            _render_image_inline(paragraph, item, template)
        elif item["type"] == "hardbreak":
            run = paragraph.add_run()
            run.add_break()


def _render_link(paragraph, content_items, base_style, link_style, template, href):
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    if not content_items:
        run = paragraph.add_run(href)
        _apply_run_style(run, {**base_style, **link_style})
        return

    for item in content_items:
        if item["type"] == "text":
            run = paragraph.add_run(item["content"])
            _apply_run_style(run, {**base_style, **link_style})
        elif item["type"] == "bold":
            bold_style = {**base_style, **link_style, **template.get("styles", {}).get("bold", {})}
            _render_inline_runs(paragraph, item.get("content", []), bold_style, template, link_style)
        elif item["type"] == "italic":
            italic_style = {**base_style, **link_style, **template.get("styles", {}).get("italic", {})}
            _render_inline_runs(paragraph, item.get("content", []), italic_style, template, link_style)


def _render_code_block(doc, node, template):
    style_dict = template.get("styles", {}).get("code_block", {})
    p = doc.add_paragraph()
    run = p.add_run(node.get("content", ""))
    _apply_run_style(run, style_dict)

    if "bg_color" in style_dict:
        _apply_paragraph_shading(p, style_dict["bg_color"])
    if "border_color" in style_dict and "border_width_pt" in style_dict:
        _apply_paragraph_border(p, style_dict["border_color"], style_dict["border_width_pt"])


def _render_table(doc, node, template):
    headers = node.get("headers", [])
    rows = node.get("rows", [])
    if not headers and not rows:
        return

    num_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if num_cols == 0:
        return

    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"

    header_style = template.get("styles", {}).get("table_header", {})
    cell_style = template.get("styles", {}).get("table_cell", {})

    for col_idx, cell_data in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        p = cell.paragraphs[0]
        if cell_data.get("inline"):
            _render_inline_runs(p, cell_data["inline"], header_style, template)
        else:
            run = p.add_run(cell_data.get("text", ""))
            _apply_run_style(run, header_style)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            p = cell.paragraphs[0]
            if cell_data.get("inline"):
                _render_inline_runs(p, cell_data["inline"], cell_style, template)
            else:
                run = p.add_run(cell_data.get("text", ""))
                _apply_run_style(run, cell_style)


def _render_list(doc, node, template):
    list_type = node.get("list_type", "bullet")
    for item in node.get("items", []):
        _render_list_item(doc, item, template, list_type, level=0)


def _render_list_item(doc, item, template, list_type, level):
    content_parts = item.get("content", [])
    for part in content_parts:
        if part["type"] == "paragraph":
            p = doc.add_paragraph()
            if list_type == "bullet":
                p.style = doc.styles["List Bullet"]
            else:
                p.style = doc.styles["List Number"]

            if level > 0:
                pf = p.paragraph_format
                pf.left_indent = Cm(1.27 * (level + 1))

            if part.get("inline"):
                _render_inline_runs(p, part["inline"], template.get("styles", {}).get("paragraph", {}), template)
            else:
                run = p.add_run(part.get("text", ""))
                _apply_run_style(run, template.get("styles", {}).get("paragraph", {}))
        elif part["type"] == "list":
            for sub_item in part.get("items", []):
                _render_list_item(doc, sub_item, template, part.get("list_type", "bullet"), level + 1)


def _render_blockquote(doc, node, template):
    for item in node.get("content", []):
        if item["type"] == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            if item.get("inline"):
                _render_inline_runs(p, item["inline"], template.get("styles", {}).get("paragraph", {}), template)
            else:
                run = p.add_run(item.get("text", ""))
                _apply_run_style(run, template.get("styles", {}).get("paragraph", {}))


def _render_image_inline(paragraph, item, template):
    pass


def convert(md_text, template=None, file_path=""):
    if template is None:
        template = deepcopy(DEFAULT_TEMPLATE)
    else:
        t = deepcopy(DEFAULT_TEMPLATE)
        _deep_merge(t, template)
        template = t

    tokens = parse_markdown(md_text)
    ir = build_ir(tokens)

    if len(ir) > MAX_IR_NODES:
        import warnings
        warnings.warn(f"IR 节点数 ({len(ir)}) 超过 {MAX_IR_NODES}，仅渲染前 {MAX_IR_NODES} 个节点")

    doc = render_docx(ir, template)

    if file_path:
        md_dir = Path(file_path).parent
        _embed_images(doc, ir, md_dir)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _deep_merge(base, override):
    for key, value in override.items():
        if key in base and isinstance(base[key], dict) and isinstance(value, dict):
            _deep_merge(base[key], value)
        else:
            base[key] = value


def _embed_images(doc, ir, md_dir):
    from PIL import Image as PILImage

    inline_shapes = doc.inline_shapes
    image_index = 0

    def _find_images_in_ir(nodes):
        nonlocal image_index
        for node in nodes:
            if node.get("inline"):
                for inline_item in node["inline"]:
                    if inline_item["type"] == "image":
                        src = inline_item.get("src", "")
                        if not src.startswith(("http://", "https://")):
                            img_path = md_dir / src
                            if img_path.exists():
                                try:
                                    img = PILImage.open(img_path)
                                    width_px, height_px = img.size
                                    for paragraph in doc.paragraphs:
                                        if paragraph.text == "" or image_index >= 2000:
                                            break
                                    image_index += 1
                                except Exception:
                                    pass
                            else:
                                for paragraph in doc.paragraphs:
                                    if _extract_plain_text_from_paragraph(paragraph) == "":
                                        run = paragraph.add_run(f"[图片未找到: {src}]")
                                        run.italic = True
                                        break
            for key in ("items", "content", "rows"):
                if key in node and isinstance(node[key], list):
                    _find_images_in_ir(node[key])
            for row in node.get("rows", []):
                for cell in row:
                    if cell.get("inline"):
                        _find_images_in_ir([cell])

    _find_images_in_ir(ir)


def _extract_plain_text_from_paragraph(paragraph):
    return paragraph.text
